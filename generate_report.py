import os
import urllib.request
import zlib
import base64
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def download_mermaid_image(mermaid_code, filename):
    compressed = zlib.compress(mermaid_code.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('ascii')
    url = f"https://kroki.io/mermaid/png/{encoded}"
    try:
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req) as response, open(filename, 'wb') as out_file:
            out_file.write(response.read())
        return True
    except Exception as e:
        print(f"Failed to download diagram {filename} from {url}: {e}")
        return False

def add_heading(doc, text, level, font_name='Times New Roman', font_size=16):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = True
    # Word sometimes ignores the font name unless set like this:
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return heading

def add_paragraph(doc, text, font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(text)
    p.alignment = alignment
    # Set 1.5 line spacing
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
    return p

def add_code_section(doc, code_text):
    p = doc.add_paragraph(code_text)
    p.paragraph_format.line_spacing = 1.0
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(11)
    return p

def set_document_font(doc, font_name='Times New Roman', font_size=12):
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def main():
    doc = Document()
    
    # Configure page setup (A4, 1-inch margins)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    
    set_document_font(doc)

    # Main Title
    title = doc.add_paragraph("Analyzing the Role of AI-Assisted Programming in Modern Software Development")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(20)
    title_run.font.bold = True

    add_paragraph(doc, "")

    # Task 1: Conceptual Understanding
    add_heading(doc, "Task 1: Conceptual Understanding", level=1)
    
    add_heading(doc, "What is AI-Assisted Programming?", level=2, font_size=14)
    add_paragraph(doc, "AI-Assisted Programming refers to the use of artificial intelligence tools and models, particularly Large Language Models (LLMs), to aid software developers in writing, debugging, testing, and optimizing code. Rather than replacing developers, these tools act as intelligent co-pilots that can understand natural language prompts and translate them into functional code blocks, significantly accelerating the software development lifecycle.")
    
    add_heading(doc, "Types of AI coding tools", level=2, font_size=14)
    add_paragraph(doc, "1. Code Generation Tools: Tools like GitHub Copilot and ChatGPT that can generate entire functions or classes based on a description.\n2. Debugging Assistants: Tools that analyze error logs or stack traces and suggest fixes.\n3. Refactoring Tools: AI plugins that suggest cleaner, more efficient ways to write existing code, such as SonarQube with AI capabilities.\n4. Documentation Generators: AI tools that automatically write docstrings and project documentation based on the source code.")

    add_heading(doc, "Advantages in software development", level=2, font_size=14)
    add_paragraph(doc, "The primary advantage is a drastic reduction in development time. Developers spend less time typing boilerplate code or searching for syntax on StackOverflow. AI also helps lower the barrier to entry for complex frameworks, improves code readability through automated refactoring, and can rapidly prototype applications to validate ideas faster.")

    add_heading(doc, "Risks and limitations", level=2, font_size=14)
    add_paragraph(doc, "Despite the benefits, AI tools present significant risks. 'Hallucinations' occur when an AI generates syntactically correct but functionally incorrect or nonsensical code. There are also security issues, as AI might inadvertently introduce vulnerabilities (like SQL injection flaws) if trained on insecure code. Dependency issues arise when AI suggests outdated or deprecated libraries. Furthermore, over-reliance on AI can degrade a developer's foundational problem-solving skills.")

    doc.add_page_break()

    # Task 2: Practical Implementation
    add_heading(doc, "Task 2: Practical Implementation", level=1)
    add_paragraph(doc, "For this task, a Library Management System was developed using Python. The problem is of moderate complexity, involving entities such as Books, Users, and Transactions.")
    
    add_heading(doc, "Comparison of Implementations", level=2, font_size=14)
    add_paragraph(doc, "Version A (Traditional Approach): Written entirely manually using procedural programming and in-memory lists saved to JSON files. It lacks advanced error handling and data integrity constraints.")
    add_paragraph(doc, "Version B (AI-Assisted Approach): Developed with the help of AI prompts. It utilizes Object-Oriented Principles and a SQLite database. It includes comprehensive logging, foreign key constraints, and robust input validation.")
    
    add_paragraph(doc, "Development Time: Version B was completed significantly faster because the AI instantly generated the SQLite table schemas and the boilerplate logic for database interaction, whereas Version A required manual implementation of data saving/loading logic.")
    add_paragraph(doc, "Code Quality & Error Handling: Version B exhibits much higher code quality. The AI naturally included try-except blocks for database IntegrityErrors, which would have been tedious to handle manually in Version A. Version A's file-based storage is prone to data corruption compared to Version B's ACID-compliant SQLite approach.")
    add_paragraph(doc, "Readability and Maintainability: Version B's Object-Oriented design makes it highly modular and easier to maintain. Version A's procedural design with global variables is harder to scale.")

    add_heading(doc, "Code Sections Identification", level=2, font_size=14)
    add_paragraph(doc, "Manually Written Code (Version A): All functions including add_book(), register_user(), borrow_book(), return_book(), and the main loop in library_system.py were manually constructed.")
    add_paragraph(doc, "AI-Generated Code (Version B): The DatabaseManager class, SQLite queries, try-except integrity blocks, and the basic skeleton of the CLI loop were generated via prompts and then refined.")

    doc.add_page_break()
    
    add_heading(doc, "UML Diagrams", level=2, font_size=14)
    add_paragraph(doc, "The following UML diagrams were generated using Mermaid.")

    # Mermaid definitions
    use_case = '''
    usecaseDiagram
    actor User
    actor Librarian
    usecase "Borrow Book" as UC1
    usecase "Return Book" as UC2
    usecase "View Books" as UC3
    usecase "Add Book" as UC4
    usecase "Register User" as UC5
    User --> UC1
    User --> UC2
    User --> UC3
    Librarian --> UC4
    Librarian --> UC5
    Librarian --> UC3
    '''
    
    class_diagram = '''
    classDiagram
    class Book {
      +int id
      +String title
      +String author
      +String isbn
      +bool is_borrowed
    }
    class User {
      +int id
      +String name
      +String email
    }
    class Transaction {
      +int id
      +int user_id
      +int book_id
      +Date borrow_date
      +Date return_date
    }
    class Library {
      +add_book()
      +register_user()
      +borrow_book()
      +return_book()
    }
    Library "1" -- "*" Book
    Library "1" -- "*" User
    Library "1" -- "*" Transaction
    '''

    activity_diagram = '''
    stateDiagram-v2
    [*] --> Start
    Start --> CheckUser
    CheckUser --> UserExists? : Is User Registered?
    UserExists? --> CheckBook : Yes
    UserExists? --> Reject : No
    CheckBook --> Available? : Is Book Available?
    Available? --> Borrow : Yes
    Available? --> Reject : No
    Borrow --> UpdateDB : Mark as Borrowed
    UpdateDB --> End
    Reject --> End
    End --> [*]
    '''

    sequence_diagram = '''
    sequenceDiagram
    actor User
    participant CLI
    participant Library
    participant Database
    User->>CLI: Request Borrow (User_ID, Book_ID)
    CLI->>Library: borrow_book(User_ID, Book_ID)
    Library->>Database: SELECT id FROM users WHERE id=User_ID
    Database-->>Library: Return User
    Library->>Database: SELECT is_borrowed FROM books WHERE id=Book_ID
    Database-->>Library: Return Status
    Library->>Database: UPDATE books SET is_borrowed=1
    Library->>Database: INSERT INTO transactions...
    Library-->>CLI: Return True (Success)
    CLI-->>User: Display "Success"
    '''

    images = [
        (use_case, "use_case.png", "Use Case Diagram"),
        (class_diagram, "class.png", "Class Diagram"),
        (activity_diagram, "activity.png", "Activity Diagram"),
        (sequence_diagram, "sequence.png", "Sequence Diagram")
    ]

    for code, filename, title in images:
        if download_mermaid_image(code, filename):
            add_paragraph(doc, f"{title}:", font_name='Times New Roman', font_size=12)
            doc.add_picture(filename, width=Inches(6.0))
            add_paragraph(doc, "")

    doc.add_page_break()

    # AI Prompts
    add_heading(doc, "AI Prompts Used", level=2, font_size=14)
    prompts = """
    Prompt 1: Initial Architecture
    "Write a Python Library Management System that uses SQLite for the database. It should have tables for books, users, and transactions. Use Object-Oriented Programming principles. The library class should have methods to add books, register users, borrow books, and return books."
    
    Prompt 2: Refinement - Error Handling
    "Enhance the previous code by adding robust error handling. Use Python's built-in logging module to log database errors (e.g., IntegrityError for duplicate ISBNs or emails). Also, ensure the borrow_book method checks if the user exists and if the book is already borrowed."
    
    Prompt 3: Refinement - CLI Interface
    "Now write a main_menu function that provides a Command Line Interface (CLI) for the library system. It should run in a loop and handle user inputs safely, catching ValueError if the user enters non-integers for IDs."
    """
    add_code_section(doc, prompts)

    doc.add_page_break()

    # Task 3: Critical Analysis
    add_heading(doc, "Task 3: Critical Analysis", level=1)
    
    add_heading(doc, "How AI changed the coding approach", level=2, font_size=14)
    add_paragraph(doc, "AI fundamentally shifted my coding approach from 'writing syntax' to 'designing architecture'. Instead of worrying about the exact SQL syntax to create tables with foreign keys, I could focus on specifying the business rules (e.g., users, books, and borrowing rules). AI acted as an accelerator, executing the low-level implementation while I managed the high-level logic.")

    add_heading(doc, "Whether AI improved or reduced understanding", level=2, font_size=14)
    add_paragraph(doc, "In this case, AI improved understanding by demonstrating best practices. For example, it naturally suggested using the logging module and creating a separate DatabaseManager class to abstract SQL queries, which is a superior design pattern. However, for a total beginner, there is a risk of reduced understanding if the code is simply copy-pasted without reviewing how the SQLite cursor objects operate.")

    add_heading(doc, "In which cases AI should NOT be used", level=2, font_size=14)
    add_paragraph(doc, "AI should not be used blindly in highly secure environments (e.g., banking software dealing with PII) unless the output is rigorously audited, as it may hallucinate insecure cryptographic practices. It should also be avoided when dealing with proprietary or highly novel algorithms where the AI's training data has no relevant context, leading to poor or misleading suggestions.")

    add_heading(doc, "Ethical considerations in AI-generated code", level=2, font_size=14)
    add_paragraph(doc, "There are concerns regarding intellectual property and copyright, as AI models are trained on vast amounts of open-source code without explicit attribution to the original authors. Furthermore, developers must take responsibility for AI-generated code; if an AI suggests code that causes a system outage or security breach, the ethical and legal burden falls on the human developer who approved it.")

    doc.add_page_break()

    # Task 4: Conclusion
    add_heading(doc, "Task 4: Conclusion", level=1)
    add_paragraph(doc, "AI in programming education and industry represents a paradigm shift comparable to the transition from assembly language to high-level languages. In the industry, it is becoming an indispensable tool for productivity, allowing teams to deliver software faster and with fewer boilerplate errors. In education, while there are fears of students using AI to bypass learning, it can actually serve as a powerful personalized tutor. Ultimately, AI will not replace software engineers; rather, engineers who leverage AI will replace those who do not.")

    add_paragraph(doc, "")
    add_heading(doc, "GitHub Repository", level=2, font_size=14)
    add_paragraph(doc, "The source code, documentation, and commit history for this assessment can be found at the following GitHub repository link:")
    add_paragraph(doc, "[INSERT YOUR GITHUB REPOSITORY LINK HERE]")
    
    doc.save("Assessment_Report.docx")
    print("Document Assessment_Report.docx generated successfully.")

if __name__ == '__main__':
    main()
